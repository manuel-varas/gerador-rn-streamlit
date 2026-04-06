import streamlit as st
import re
from datetime import date
import os
import json
import urllib.request
import copy
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# =============================
# CONFIG
# =============================
st.set_page_config(page_title="Gerador RN - Allianz", layout="wide")
st.title("Gerador de RN - Modelo Word")
st.success("✅ App carregado com sucesso")

TEMPLATE = "MODELO RN (1).docx"

# =============================
# NORMALIZAÇÃO (SEM norm() — SÓ _norm())
# =============================
def _norm(s: str) -> str:
    s = (s or "").replace("\u2019", "'").replace("\u2018", "'")
    return re.sub(r"\s+", " ", s).strip()

# =============================
# STATE (robusto / incremental)
# =============================
# ✅ começa com 1 local
if "n_locais" not in st.session_state:
    st.session_state.n_locais = 1

if "locais_version" not in st.session_state:
    st.session_state.locais_version = 0

if "locais_data" not in st.session_state:
    st.session_state.locais_data = [
        {"cep": "", "endereco_base": "", "numero": "", "complemento": "", "atividade": ""}
        for _ in range(st.session_state.n_locais)
    ]

if "vr_data" not in st.session_state:
    st.session_state.vr_data = [
        {"predio": "R$ ", "mmu": "R$ ", "mmp": "R$ ", "lucros": "R$ "}
        for _ in range(st.session_state.n_locais)
    ]

# Página 2: I e III
for k in ["segurado_p2", "cnpj_p2", "cossegurados", "cosseg_cnpj", "atividade_principal"]:
    st.session_state.setdefault(k, "")

# Vigência
st.session_state.setdefault("vig_inicio", date.today())
st.session_state.setdefault("vig_fim", date.today())

# Página 1
for k, default in {
    "rn": "",
    "destinatario": "",
    "subscritor": "",
    "filial": "",
    "email_user": "",
    "cotacao": "Riscos Nomeados",
    "segurado_p1": "",
    "cnpj_p1": "",
    "paginas": 13,
}.items():
    st.session_state.setdefault(k, default)

st.session_state.setdefault("data_doc", date.today())
st.session_state.setdefault("generated_docx_bytes", None)

# =============================
# COSSEGURO (Página 9)
# =============================
if "cosseguro_data" not in st.session_state:
    st.session_state.cosseguro_data = [
        {"seguradora": "Allianz Seguros S.A.", "susep": "05177", "pct": "100,00%", "lmi": "R$ "},
        {"seguradora": "", "susep": "", "pct": "", "lmi": "R$ "},
    ]

# =============================
# IX - LMGA (Página 9)
# =============================
if "lmga_data" not in st.session_state:
    st.session_state.lmga_data = []  # será carregado do template

# =============================
# HELPERS
# =============================
def safe_rerun():
    try:
        st.rerun()
    except Exception:
        st.experimental_rerun()

def _sync_lists():
    """Garante que Locais e VR cresçam juntos (sempre)."""
    n = int(st.session_state.n_locais)

    L = st.session_state.locais_data
    if len(L) < n:
        L.extend([{"cep": "", "endereco_base": "", "numero": "", "complemento": "", "atividade": ""} for _ in range(n - len(L))])
    elif len(L) > n:
        st.session_state.locais_data = L[:n]

    V = st.session_state.vr_data
    if len(V) < n:
        V.extend([{"predio": "R$ ", "mmu": "R$ ", "mmp": "R$ ", "lucros": "R$ "} for _ in range(n - len(V))])
    elif len(V) > n:
        st.session_state.vr_data = V[:n]

def aumentar_locais(mais=1):
    st.session_state.n_locais = int(st.session_state.n_locais) + int(mais)
    _sync_lists()

def reduzir_locais(menos=1):
    # ✅ mínimo agora é 1
    st.session_state.n_locais = max(1, int(st.session_state.n_locais) - int(menos))
    _sync_lists()

# =============================
# FORMAT HELPERS (R$ e %)
# =============================
def format_cnpj(cnpj: str) -> str:
    nums = re.sub(r"\D", "", cnpj or "")
    if len(nums) == 14:
        return f"{nums[:2]}.{nums[2:5]}.{nums[5:8]}/{nums[8:12]}-{nums[12:]}"
    return cnpj

def format_cep(cep: str) -> str:
    nums = re.sub(r"\D", "", cep or "")
    if len(nums) == 8:
        return f"{nums[:5]}-{nums[5:]}"
    return cep

def viacep_lookup(cep: str) -> str:
    nums = re.sub(r"\D", "", cep or "")
    if len(nums) != 8:
        return ""
    url = f"https://viacep.com.br/ws/{nums}/json/"
    try:
        with urllib.request.urlopen(url, timeout=6) as resp:
            data = json.loads(resp.read().decode("utf-8"))
        if data.get("erro"):
            return ""
        logradouro = data.get("logradouro", "")
        bairro = data.get("bairro", "")
        localidade = data.get("localidade", "")
        uf = data.get("uf", "")
        complemento = data.get("complemento", "")
        parts = [p for p in [logradouro, complemento, bairro, f"{localidade}-{uf}"] if p]
        return " - ".join(parts).strip()
    except Exception:
        return ""

def montar_endereco_final(endereco_base: str, numero: str, complemento: str) -> str:
    parts = []
    base = (endereco_base or "").strip()
    if base:
        parts.append(base)
    num = (numero or "").strip()
    if num:
        parts.append(f"Nº {num}")
    comp = (complemento or "").strip()
    if comp:
        parts.append(comp)
    return " - ".join(parts)

def parse_brl_number(val: str) -> float:
    if val is None:
        return 0.0
    s = str(val).strip()
    if not s:
        return 0.0

    s = s.replace("R$", "").replace("r$", "").replace(" ", "")
    if not s:
        return 0.0

    if "," in s and "." in s:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    else:
        if "," in s:
            s = s.replace(".", "").replace(",", ".")

    try:
        return float(s)
    except Exception:
        return 0.0

def fmt_brl_number(x: float) -> str:
    s = f"{x:,.2f}"
    return s.replace(",", "X").replace(".", ",").replace("X", ".")

def fmt_brl_money(x: float) -> str:
    return f"R$ {fmt_brl_number(x)}"

def ensure_prefix(v: str) -> str:
    txt = (v or "").strip()
    if txt.startswith("R$"):
        return "R$ " if txt == "R$" else txt
    if txt == "":
        return "R$ "
    return "R$ " + txt

def format_money_field(key: str):
    raw = ensure_prefix(st.session_state.get(key, ""))
    value = parse_brl_number(raw)
    if raw.strip() in ("R$", ""):
        st.session_state[key] = "R$ "
    else:
        st.session_state[key] = fmt_brl_money(value)

def parse_percent(val: str) -> float:
    if val is None:
        return 0.0
    s = str(val).strip().replace(" ", "")
    if not s:
        return 0.0
    s = s.replace("%", "")

    if "," in s and "." in s:
        if s.rfind(",") > s.rfind("."):
            s = s.replace(".", "").replace(",", ".")
        else:
            s = s.replace(",", "")
    else:
        if "," in s:
            s = s.replace(".", "").replace(",", ".")

    try:
        x = float(s)
    except Exception:
        return 0.0

    if 0 < x <= 1:
        return x * 100.0
    return x

def fmt_percent_br(p: float) -> str:
    s = f"{p:,.2f}"
    s = s.replace(",", "X").replace(".", ",").replace("X", ".")
    return f"{s}%"

def format_percent_field(key: str):
    raw = st.session_state.get(key, "")
    p = parse_percent(raw)
    st.session_state[key] = "" if p <= 0 else fmt_percent_br(p)

# =============================
# WORD HELPERS
# =============================
def set_cell_text(cell, text, paragraph_index=0):
    """Escreve preservando estilo (não destrói a célula)."""
    while len(cell.paragraphs) <= paragraph_index:
        cell.add_paragraph("")
    p = cell.paragraphs[paragraph_index]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)

def clear_cell_keep_format(cell):
    tc = cell._tc
    for p in list(tc.p_lst):
        tc.remove(p)

def replace_in_cell_all(cell, old, new):
    for p in cell.paragraphs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)

def find_table(doc, anchor_text):
    a = (anchor_text or "").upper()
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                if a in (c.text or "").upper():
                    return t
    return None

def find_row(table, left_label_contains):
    needle = (left_label_contains or "").upper()
    for i, row in enumerate(table.rows):
        if len(row.cells) >= 2 and needle in (row.cells[0].text or "").upper():
            return i
    return None

def find_locais_table(doc):
    for t in doc.tables:
        if len(t.rows) == 0:
            continue
        header = " ".join((c.text or "").strip().upper() for c in t.rows[0].cells)
        if ("LOCAL" in header) and ("ENDEREÇO" in header) and ("ATIVIDADE" in header) and len(t.columns) >= 3:
            return t
    return None

def find_vr_table(doc):
    for t in doc.tables:
        if len(t.rows) < 2:
            continue
        header = " ".join((c.text or "").strip().upper() for c in t.rows[1].cells)
        if ("PRÉDIO" in header) and ("MMU" in header) and ("MMP" in header) and ("LUCROS" in header) and ("LOCAL" in header):
            return t
    return None

def ensure_table_rows_with_style(table, desired_data_rows, header_rows=1, template_row_index=None):
    current_rows = len(table.rows)
    target_rows = header_rows + desired_data_rows
    if current_rows >= target_rows:
        return
    if template_row_index is None:
        template_row_index = current_rows - 1
    template_tr = table.rows[template_row_index]._tr
    for _ in range(target_rows - current_rows):
        new_tr = copy.deepcopy(template_tr)
        table._tbl.append(new_tr)

def vr_adjust_rows(table, desired_rows):
    totals_idx = None
    for i, row in enumerate(table.rows):
        if any("TOTAIS" in (c.text or "").upper() for c in row.cells):
            totals_idx = i
            break
    if totals_idx is None:
        return

    data_start = 2
    current_data = totals_idx - data_start

    if desired_rows > current_data:
        template_tr = table.rows[totals_idx - 1]._tr
        for _ in range(desired_rows - current_data):
            new_tr = copy.deepcopy(template_tr)
            table._tbl.insert(totals_idx, new_tr)
            totals_idx += 1

    elif desired_rows < current_data:
        for _ in range(current_data - desired_rows):
            remove_idx = totals_idx - 1
            table._tbl.remove(table.rows[remove_idx]._tr)
            totals_idx -= 1

# =============================
# VI - COBERTURAS (do modelo)
# =============================
def find_vi_table(doc: Document):
    for t in doc.tables:
        if len(t.rows) < 2:
            continue
        header = " ".join(_norm(c.text).upper() for c in t.rows[0].cells)
        if ("LOCAIS" in header) and ("GARANTIAS" in header) and ("LMI" in header) and ("FRANQUIA" in header or "POS" in header):
            return t
    return None

def extract_vi_from_template():
    doc = Document(TEMPLATE)
    t = find_vi_table(doc)
    if not t:
        return []
    items = []
    secao = ""
    for r in t.rows[1:]:
        alltxt = " ".join(_norm(c.text) for c in r.cells)
        if "GARANTIA BÁSICA" in alltxt.upper():
            secao = "Garantia Básica"
            continue
        if "GARANTIAS ADICIONAIS" in alltxt.upper():
            secao = "Garantias Adicionais"
            continue
        if len(r.cells) >= 4:
            loc = _norm(r.cells[0].text)
            gar = _norm(r.cells[1].text)
            pos = _norm(r.cells[3].text)
            if loc and gar:
                items.append({
                    "secao": secao or "Coberturas",
                    "locais": loc,
                    "include": False,   # ✅ default desmarcado
                    "garantia": gar,
                    "lmi": "R$ ",
                    "pos": pos
                })
    return items

if "coberturas_data" not in st.session_state:
    st.session_state.coberturas_data = extract_vi_from_template()

def fill_vi_in_word(doc: Document):
    """
    ✅ Regra:
    - include == False  -> remove a linha da tabela VI no Word (não aparece)
    - include == True   -> preenche LMI / POS (e mantém linha)
    """
    t = find_vi_table(doc)
    if not t:
        return

    cov_map = {}
    for item in st.session_state.coberturas_data:
        key = _norm(item.get("garantia"))
        if key:
            cov_map[key] = item

    rows_to_remove = []

    for ridx in range(1, len(t.rows)):
        r = t.rows[ridx]
        if len(r.cells) < 4:
            continue

        gar = _norm(r.cells[1].text)
        if not gar:
            continue

        upper = gar.upper()
        if "GARANTIA BÁSICA" in upper or "GARANTIAS ADICIONAIS" in upper:
            continue

        if gar in cov_map:
            item = cov_map[gar]

            if not item.get("include", False):
                rows_to_remove.append(ridx)
                continue

            lmi = item.get("lmi", "R$ ")
            pos = item.get("pos", "")

            set_cell_text(r.cells[2], (lmi if (lmi and lmi.strip() != "R$") else ""), 0)
            set_cell_text(r.cells[3], pos or "", 0)

    for ridx in sorted(rows_to_remove, reverse=True):
        t._tbl.remove(t.rows[ridx]._tr)

# =============================
# IX - LMGA (Cobertura | Limite (R$))
# =============================
def find_lmga_table(doc: Document):
    candidates = []
    for t in doc.tables:
        if len(t.rows) < 2:
            continue
        header = " ".join(_norm(c.text).upper() for c in t.rows[0].cells)
        if ("COBERTURA" in header) and ("LIMITE" in header):
            candidates.append(t)
    if not candidates:
        return None
    for t in candidates:
        alltxt = " ".join(_norm(c.text).upper() for row in t.rows for c in row.cells)
        if "TOTAL" in alltxt:
            return t
    return max(candidates, key=lambda x: len(x.rows))

def extract_lmga_from_template():
    doc = Document(TEMPLATE)
    t = find_lmga_table(doc)
    if not t:
        return []
    items = []
    for r in t.rows[1:]:
        if len(r.cells) < 2:
            continue
        cov = _norm(r.cells[0].text)
        if not cov:
            continue
        if "TOTAL" in cov.upper():
            break
        items.append({"cobertura": cov, "limite": "R$ "})
    return items

if not st.session_state.lmga_data:
    st.session_state.lmga_data = extract_lmga_from_template()

def fill_lmga_in_word(doc: Document):
    t = find_lmga_table(doc)
    if not t:
        return
    mp = {_norm(d["cobertura"]): d for d in st.session_state.lmga_data if d.get("cobertura")}
    for r in t.rows[1:]:
        if len(r.cells) < 2:
            continue
        cov = _norm(r.cells[0].text)
        if not cov:
            continue
        if "TOTAL" in cov.upper():
            break
        if cov in mp:
            lim = mp[cov].get("limite", "R$ ")
            set_cell_text(r.cells[1], (lim if (lim and lim.strip() != "R$") else ""), 0)

# =============================
# COSSEGURO - Word
# =============================
def find_cosseguro_table(doc: Document):
    for t in doc.tables:
        if len(t.rows) < 2:
            continue
        header = " ".join((c.text or "").strip().upper() for c in t.rows[0].cells)
        if ("SEGURADORA" in header) and ("SUSEP" in header) and ("PARTICIPA" in header) and ("LMI" in header):
            return t
    return None

def cosseguro_adjust_rows(table, desired_rows):
    total_idx = None
    for i, row in enumerate(table.rows):
        row_txt = " ".join((c.text or "").strip().upper() for c in row.cells)
        if "TOTAL" in row_txt:
            total_idx = i
            break
    if total_idx is None:
        return

    header_rows = 1
    current_data = total_idx - header_rows

    if desired_rows > current_data:
        template_tr = table.rows[total_idx - 1]._tr
        for _ in range(desired_rows - current_data):
            new_tr = copy.deepcopy(template_tr)
            table._tbl.insert(total_idx, new_tr)
            total_idx += 1

    elif desired_rows < current_data:
        for _ in range(current_data - desired_rows):
            remove_idx = total_idx - 1
            table._tbl.remove(table.rows[remove_idx]._tr)
            total_idx -= 1

def fill_cosseguro_in_word(doc: Document):
    t = find_cosseguro_table(doc)
    if not t:
        return

    data = st.session_state.cosseguro_data
    desired = len(data)
    cosseguro_adjust_rows(t, desired)

    total_idx = None
    for i, row in enumerate(t.rows):
        row_txt = " ".join((c.text or "").strip().upper() for c in row.cells)
        if "TOTAL" in row_txt:
            total_idx = i
            break
    if total_idx is None:
        return

    for i in range(desired):
        row_idx = 1 + i
        if row_idx >= total_idx:
            break

        seg = data[i].get("seguradora", "")
        susep = data[i].get("susep", "")
        pct = data[i].get("pct", "")
        lmi = data[i].get("lmi", "")

        set_cell_text(t.cell(row_idx, 0), seg)
        set_cell_text(t.cell(row_idx, 1), susep)
        set_cell_text(t.cell(row_idx, 2), pct)
        set_cell_text(t.cell(row_idx, 3), (lmi if (lmi and lmi.strip() != "R$") else ""))

# =============================
# GERAR WORD (mantém seu fluxo)
# =============================
def build_docx_bytes():
    doc = Document(TEMPLATE)
    n = int(st.session_state.n_locais)

    # capa
    cover = find_table(doc, "PROC. Nº")
    if cover:
        i = find_row(cover, "PROC. Nº")
        if i is not None and st.session_state.rn:
            set_cell_text(cover.cell(i, 1), f"RN - {st.session_state.rn}")

        i = find_row(cover, "DESTINATÁRIO")
        if i is not None and st.session_state.destinatario:
            set_cell_text(cover.cell(i, 1), st.session_state.destinatario)

        i = find_row(cover, "REMETENTE/FROM")
        if i is not None:
            if st.session_state.subscritor:
                set_cell_text(cover.cell(i, 1), st.session_state.subscritor, 0)
            if st.session_state.filial:
                set_cell_text(cover.cell(i, 1), st.session_state.filial, 1)

        i = find_row(cover, "DEPTO/DIVISION")
        if i is not None and st.session_state.email_user:
            replace_in_cell_all(cover.cell(i, 1), "xxxx.xxxx", st.session_state.email_user)

        i = find_row(cover, "DATA/DATE")
        if i is not None:
            set_cell_text(cover.cell(i, 1), st.session_state.data_doc.strftime("%d/%m/%Y"))

        i = find_row(cover, "PÁGINAS/PAGES")
        if i is not None:
            set_cell_text(cover.cell(i, 1), f"{int(st.session_state.paginas)} (incluindo esta capa/including the cover page)")

    # cotação
    quote = find_table(doc, "COTAÇÃO:")
    if quote:
        i = find_row(quote, "COTAÇÃO")
        if i is not None:
            set_cell_text(quote.cell(i, 1), st.session_state.cotacao)

        i = find_row(quote, "SEGURADO")
        if i is not None and st.session_state.segurado_p1:
            set_cell_text(quote.cell(i, 1), st.session_state.segurado_p1)

        i = find_row(quote, "CNPJ")
        if i is not None and st.session_state.cnpj_p1:
            set_cell_text(quote.cell(i, 1), format_cnpj(st.session_state.cnpj_p1))

    # pág 2 segurado/cosseg
    t_seg = find_table(doc, "I – Segurado")
    if t_seg and len(t_seg.rows) >= 4 and len(t_seg.columns) >= 2:
        set_cell_text(t_seg.cell(1, 0), st.session_state.segurado_p2)
        set_cell_text(t_seg.cell(1, 1), format_cnpj(st.session_state.cnpj_p2))
        set_cell_text(t_seg.cell(3, 0), st.session_state.cossegurados)
        set_cell_text(t_seg.cell(3, 1), format_cnpj(st.session_state.cosseg_cnpj))

    # atividade principal
    t_iii = find_table(doc, "III – Objeto Segurado / Atividade Principal")
    if t_iii and len(t_iii.rows) >= 5:
        set_cell_text(t_iii.cell(4, 0), st.session_state.atividade_principal)

    # vigência
    t_vig = find_table(doc, "IV – Vigência do seguro")
    if t_vig and len(t_vig.rows) >= 2 and len(t_vig.columns) >= 2:
        cell = t_vig.cell(1, 1)
        clear_cell_keep_format(cell)
        p1 = cell.add_paragraph(f"Das 24 horas do dia {st.session_state.vig_inicio.strftime('%d/%m/%Y')}")
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2 = cell.add_paragraph(f"Às 24 horas do dia {st.session_state.vig_fim.strftime('%d/%m/%Y')}")
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ✅ LOCAIS: só mantém linhas com endereço (remove vazias no Word)
    t_locais = find_locais_table(doc)
    if t_locais:
        valid_idx = []
        for i in range(n):
            end_base = (st.session_state.locais_data[i].get("endereco_base") or "").strip()
            if end_base:
                valid_idx.append(i)

        ensure_table_rows_with_style(t_locais, desired_data_rows=len(valid_idx), header_rows=1)

        for j, i in enumerate(valid_idx):
            row_index = 1 + j
            local_num = f"{j+1:02d}"

            end_base = (st.session_state.locais_data[i].get("endereco_base") or "").strip()
            numv = (st.session_state.locais_data[i].get("numero") or "").strip()
            comp = (st.session_state.locais_data[i].get("complemento") or "").strip()
            atv = (st.session_state.locais_data[i].get("atividade") or "").strip()

            endereco_final = montar_endereco_final(end_base, numv, comp)

            set_cell_text(t_locais.cell(row_index, 0), local_num)
            set_cell_text(t_locais.cell(row_index, 1), endereco_final)
            set_cell_text(t_locais.cell(row_index, 2), atv)

    # VR (mantém seu fluxo original com n_locais)
    t_vr = find_vr_table(doc)
    if t_vr:
        vr_adjust_rows(t_vr, n)

        totals_idx = None
        for idx, row in enumerate(t_vr.rows):
            if any("TOTAIS" in (c.text or "").upper() for c in row.cells):
                totals_idx = idx
                break

        data_start = 2
        total_pred = total_mmu = total_mmp = total_dm = total_luc = 0.0

        for i in range(n):
            row_idx = data_start + i
            local_num = f"{i+1:02d}"

            pred = parse_brl_number(st.session_state.vr_data[i].get("predio", "R$ "))
            mmu  = parse_brl_number(st.session_state.vr_data[i].get("mmu", "R$ "))
            mmp  = parse_brl_number(st.session_state.vr_data[i].get("mmp", "R$ "))
            luc  = parse_brl_number(st.session_state.vr_data[i].get("lucros", "R$ "))
            dm = pred + mmu + mmp

            total_pred += pred
            total_mmu += mmu
            total_mmp += mmp
            total_dm += dm
            total_luc += luc

            set_cell_text(t_vr.cell(row_idx, 0), local_num)
            set_cell_text(t_vr.cell(row_idx, 1), fmt_brl_money(pred))
            set_cell_text(t_vr.cell(row_idx, 2), fmt_brl_money(mmu))
            set_cell_text(t_vr.cell(row_idx, 3), fmt_brl_money(mmp))
            set_cell_text(t_vr.cell(row_idx, 4), fmt_brl_money(dm))
            set_cell_text(t_vr.cell(row_idx, 5), fmt_brl_money(luc))

        if totals_idx is not None:
            set_cell_text(t_vr.cell(totals_idx, 1), fmt_brl_money(total_pred))
            set_cell_text(t_vr.cell(totals_idx, 2), fmt_brl_money(total_mmu))
            set_cell_text(t_vr.cell(totals_idx, 3), fmt_brl_money(total_mmp))
            set_cell_text(t_vr.cell(totals_idx, 4), fmt_brl_money(total_dm))
            set_cell_text(t_vr.cell(totals_idx, 5), fmt_brl_money(total_luc))

            vr_total = total_dm + total_luc
            vr_total_row = totals_idx + 1
            try:
                set_cell_text(t_vr.cell(vr_total_row, 5), fmt_brl_money(vr_total))
            except Exception:
                try:
                    set_cell_text(t_vr.cell(vr_total_row, 4), fmt_brl_money(vr_total))
                except Exception:
                    pass

    # VI / IX / Cosseguro
    fill_vi_in_word(doc)
    fill_lmga_in_word(doc)
    fill_cosseguro_in_word(doc)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =============================
# LIMPAR DADOS
# =============================
def limpar_dados():
    st.session_state.clear()
    safe_rerun()

# =============================
# UI TABS
# =============================
if not os.path.exists(TEMPLATE):
    st.error(f"Arquivo '{TEMPLATE}' não encontrado no repositório.")
    st.stop()

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "Página 1 - Capa/Cotação",
    "Página 2 - Segurado/Vigência/Locais",
    "Página 3 - Valor em Risco (R$)",
    "Páginas 4–8 - Coberturas (VI)",
    "Página 9 - LMGA + Cosseguro"
])

_sync_lists()

# -------- Página 1 --------
with tab1:
    c1, c2 = st.columns(2)
    with c1:
        st.text_input("PROC. Nº (RN)", key="rn")
        st.text_input("DESTINATÁRIO / To", key="destinatario")
        st.text_input("REMETENTE - Subscritor", key="subscritor")
        st.text_input("REMETENTE - Comercial / Filial", key="filial")
        st.text_input("SEGURADO (Página 1)", key="segurado_p1")
        st.text_input("CNPJ (Página 1)", key="cnpj_p1")
    with c2:
        st.text_input("E-mail (antes do @allianz.com.br)", key="email_user")
        st.date_input("DATA / DATE", key="data_doc")
        st.number_input("PÁGINAS / PAGES", min_value=1, key="paginas")
        st.text_input("COTAÇÃO", key="cotacao")

# -------- Página 2 --------
with tab2:
    st.subheader("I - Segurado / Cossegurados")
    c1, c2 = st.columns(2)
    with c1:
        st.text_input("Segurado", key="segurado_p2")
        st.text_input("Cossegurados", key="cossegurados")
    with c2:
        st.text_input("CNPJ Segurado", key="cnpj_p2")
        st.text_input("CNPJ Cossegurados", key="cosseg_cnpj")

    st.subheader("III - Atividade Principal")
    st.text_input("Atividade Principal", key="atividade_principal")

    st.subheader("IV - Vigência do seguro")
    v1, v2 = st.columns(2)
    with v1:
        st.date_input("Início de vigência", key="vig_inicio")
    with v2:
        st.date_input("Término de vigência", key="vig_fim")

    st.subheader("V - Locais em Risco/VR")
    b1, b2, b3, b4, b5 = st.columns([1, 1, 1, 1, 2])
    with b1:
        if st.button("➕ +1 local", key="btn_locais_plus_1"):
            aumentar_locais(1)
            safe_rerun()
    with b2:
        if st.button("➕ +10 locais", key="btn_locais_plus_10"):
            aumentar_locais(10)
            safe_rerun()
    with b3:
        if st.button("➖ -1 local", key="btn_locais_minus_1"):
            reduzir_locais(1)
            safe_rerun()
    with b4:
        if st.button("➖ -10 locais", key="btn_locais_minus_10"):
            reduzir_locais(10)
            safe_rerun()
    with b5:
        st.caption(f"Total de locais (Locais/VR): {st.session_state.n_locais}")

    _sync_lists()
    ver = int(st.session_state.locais_version)

    h1, h2, h3, h4, h5, h6, h7 = st.columns([0.6, 1.0, 0.9, 2.3, 0.8, 1.2, 2.0])
    h1.markdown("**Local**")
    h2.markdown("**CEP**")
    h3.markdown("**Buscar**")
    h4.markdown("**Endereço**")
    h5.markdown("**Nº**")
    h6.markdown("**Complemento**")
    h7.markdown("**Atividade**")

    for i in range(int(st.session_state.n_locais)):
        row = st.session_state.locais_data[i]
        c_local, c_cep, c_btn, c_end, c_num, c_comp, c_atv = st.columns([0.6, 1.0, 0.9, 2.3, 0.8, 1.2, 2.0])
        c_local.write(f"{i+1:02d}")

        cep_key = f"cep_{i}_{ver}"
        end_key = f"end_{i}_{ver}"
        num_key = f"num_{i}_{ver}"
        comp_key = f"comp_{i}_{ver}"
        atv_key = f"atv_{i}_{ver}"

        cep_val = c_cep.text_input("", value=row.get("cep", ""), key=cep_key, placeholder="00000-000")
        end_val = c_end.text_input("", value=row.get("endereco_base", ""), key=end_key, placeholder="Rua..., Bairro..., Cidade-UF")
        num_val = c_num.text_input("", value=row.get("numero", ""), key=num_key, placeholder="XX")
        comp_val = c_comp.text_input("", value=row.get("complemento", ""), key=comp_key, placeholder="Complemento")
        atv_val = c_atv.text_input("", value=row.get("atividade", ""), key=atv_key, placeholder="Atividade")

        st.session_state.locais_data[i]["cep"] = cep_val
        st.session_state.locais_data[i]["endereco_base"] = end_val
        st.session_state.locais_data[i]["numero"] = num_val
        st.session_state.locais_data[i]["complemento"] = comp_val
        st.session_state.locais_data[i]["atividade"] = atv_val

        if c_btn.button("CEP", key=f"buscar_{i}_{ver}"):
            base = viacep_lookup(cep_val)
            if base:
                st.session_state.locais_data[i]["cep"] = format_cep(cep_val)
                st.session_state.locais_data[i]["endereco_base"] = base
                st.session_state.locais_version += 1
                st.toast(f"CEP {format_cep(cep_val)} encontrado!", icon="✅")
                safe_rerun()
            else:
                st.toast("CEP não encontrado ou sem acesso. Preencha manualmente.", icon="⚠️")

        endereco_final_preview = montar_endereco_final(
            st.session_state.locais_data[i]["endereco_base"],
            st.session_state.locais_data[i]["numero"],
            st.session_state.locais_data[i]["complemento"],
        )
        c_end.caption(endereco_final_preview if endereco_final_preview else "")

# -------- Página 3 --------
with tab3:
    st.subheader("Valor em Risco (R$)")
    st.caption("Total DM = Prédio + MMU + MMP | VR Total = DM + Lucros")

    b1, b2, b3, b4, b5 = st.columns([1, 1, 1, 1, 2])
    with b1:
        if st.button("➕ +1 linha VR", key="btn_vr_plus_1"):
            aumentar_locais(1)
            safe_rerun()
    with b2:
        if st.button("➕ +10 linhas VR", key="btn_vr_plus_10"):
            aumentar_locais(10)
            safe_rerun()
    with b3:
        if st.button("➖ -1 linha VR", key="btn_vr_minus_1"):
            reduzir_locais(1)
            safe_rerun()
    with b4:
        if st.button("➖ -10 linhas VR", key="btn_vr_minus_10"):
            reduzir_locais(10)
            safe_rerun()
    with b5:
        st.caption(f"Total de linhas (Locais/VR): {st.session_state.n_locais}")

    _sync_lists()
    n = int(st.session_state.n_locais)

    c0, c1, c2, c3, c4, c5 = st.columns([0.6, 1.3, 1.3, 1.3, 1.4, 1.5])
    c0.markdown("**Local**")
    c1.markdown("**Prédio**")
    c2.markdown("**MMU**")
    c3.markdown("**MMP**")
    c4.markdown("**Total DM**")
    c5.markdown("**Lucros**")

    total_pred = total_mmu = total_mmp = total_dm = total_luc = 0.0

    for i in range(n):
        row = st.session_state.vr_data[i]
        r0, r1, r2, r3, r4, r5 = st.columns([0.6, 1.3, 1.3, 1.3, 1.4, 1.5])
        r0.write(f"{i+1:02d}")

        pred_key = f"vr_pred_{i}"
        mmu_key  = f"vr_mmu_{i}"
        mmp_key  = f"vr_mmp_{i}"
        luc_key  = f"vr_luc_{i}"

        st.session_state.setdefault(pred_key, row.get("predio", "R$ "))
        st.session_state.setdefault(mmu_key, row.get("mmu", "R$ "))
        st.session_state.setdefault(mmp_key, row.get("mmp", "R$ "))
        st.session_state.setdefault(luc_key, row.get("lucros", "R$ "))

        pred_s = r1.text_input("", key=pred_key, placeholder="R$ 0,00", on_change=format_money_field, args=(pred_key,))
        mmu_s  = r2.text_input("", key=mmu_key, placeholder="R$ 0,00", on_change=format_money_field, args=(mmu_key,))
        mmp_s  = r3.text_input("", key=mmp_key, placeholder="R$ 0,00", on_change=format_money_field, args=(mmp_key,))
        luc_s  = r5.text_input("", key=luc_key, placeholder="R$ 0,00", on_change=format_money_field, args=(luc_key,))

        pred = parse_brl_number(pred_s)
        mmu  = parse_brl_number(mmu_s)
        mmp  = parse_brl_number(mmp_s)
        luc  = parse_brl_number(luc_s)
        dm = pred + mmu + mmp

        st.session_state.vr_data[i]["predio"] = ensure_prefix(pred_s)
        st.session_state.vr_data[i]["mmu"] = ensure_prefix(mmu_s)
        st.session_state.vr_data[i]["mmp"] = ensure_prefix(mmp_s)
        st.session_state.vr_data[i]["lucros"] = ensure_prefix(luc_s)

        r4.write(fmt_brl_money(dm))

        total_pred += pred
        total_mmu += mmu
        total_mmp += mmp
        total_dm += dm
        total_luc += luc

    st.markdown("---")
    t0, t1, t2, t3, t4, t5 = st.columns([0.6, 1.3, 1.3, 1.3, 1.4, 1.5])
    t0.markdown("**Totais**")
    t1.markdown(f"**{fmt_brl_money(total_pred)}**")
    t2.markdown(f"**{fmt_brl_money(total_mmu)}**")
    t3.markdown(f"**{fmt_brl_money(total_mmp)}**")
    t4.markdown(f"**{fmt_brl_money(total_dm)}**")
    t5.markdown(f"**{fmt_brl_money(total_luc)}**")
    vr_total = total_dm + total_luc
    st.markdown(f"### Valor em Risco Total (DM + Lucros) = **{fmt_brl_money(vr_total)}**")

# -------- VI Coberturas --------
with tab4:
    st.subheader("VI - Coberturas, Limites e Franquias por Evento e Local")
    st.caption("✅ Marque em 'Incluir' apenas o que foi contratado. As coberturas NÃO marcadas NÃO aparecerão no Word.")

    if not st.session_state.coberturas_data:
        st.error("Não encontrei a tabela VI no modelo.")
    else:
        secao_atual = None

        for i, item in enumerate(st.session_state.coberturas_data):
            if item.get("secao") != secao_atual:
                secao_atual = item.get("secao")
                st.markdown(f"### {secao_atual}")

            colA, colG, colLMI, colPOS = st.columns([1.0, 3.4, 1.2, 2.2])

            inc_key = f"vi_inc_{i}"
            st.session_state.setdefault(inc_key, item.get("include", False))
            incluir = colA.checkbox("Incluir", key=inc_key)

            colG.text_area("Garantias", value=item.get("garantia", ""), disabled=True, height=55, key=f"vi_gar_{i}")

            lmi_key = f"vi_lmi_{i}"
            st.session_state.setdefault(lmi_key, item.get("lmi", "R$ "))
            colLMI.text_input("LMI (R$)", key=lmi_key, on_change=format_money_field, args=(lmi_key,))

            pos_key = f"vi_pos_{i}"
            st.session_state.setdefault(pos_key, item.get("pos", ""))
            colPOS.text_area("POS/Franquia (R$ / Por Evento)", key=pos_key, height=55)

            item["include"] = incluir
            item["lmi"] = st.session_state[lmi_key]
            item["pos"] = st.session_state[pos_key]

# -------- Página 9 (LMGA + Cosseguro) --------
with tab5:
    st.subheader("IX – Limite Máximo de Garantia da Apólice (LMGA)")
    st.caption("Preencha o Limite (R$) com o mesmo padrão de moeda das outras páginas.")

    if not st.session_state.lmga_data:
        st.warning("Não encontrei a tabela LMGA no modelo.")
    else:
        for i, row in enumerate(st.session_state.lmga_data):
            c1, c2 = st.columns([3.6, 1.4])
            c1.text_input("Cobertura", value=row["cobertura"], disabled=True, key=f"lmga_cov_{i}")

            lim_key = f"lmga_lim_{i}"
            st.session_state.setdefault(lim_key, row.get("limite", "R$ "))
            c2.text_input("Limite (R$)", key=lim_key, on_change=format_money_field, args=(lim_key,))
            row["limite"] = st.session_state[lim_key]

    st.markdown("---")
    st.subheader("Distribuição de Cosseguro")
    st.caption("Participação (%) formata como % e LMI – R$ como moeda.")

    cbtn1, cbtn2 = st.columns([1, 3])
    with cbtn1:
        if st.button("➕ +1 linha", key="btn_cosseguro_add"):
            st.session_state.cosseguro_data.append({"seguradora": "", "susep": "", "pct": "", "lmi": "R$ "})
            safe_rerun()
    with cbtn2:
        st.caption(f"Linhas: {len(st.session_state.cosseguro_data)}")

    h1, h2, h3, h4 = st.columns([2.5, 1.2, 1.2, 1.4])
    h1.markdown("**SEGURADORA**")
    h2.markdown("**Código SUSEP**")
    h3.markdown("**Participação (%)**")
    h4.markdown("**LMI – R$**")

    soma_pct = 0.0
    soma_lmi = 0.0

    for i, row in enumerate(st.session_state.cosseguro_data):
        c1, c2, c3, c4 = st.columns([2.5, 1.2, 1.2, 1.4])

        seg_key = f"cos_seg_{i}"
        sus_key = f"cos_sus_{i}"
        pct_key = f"cos_pct_{i}"
        lmi_key = f"cos_lmi_{i}"

        st.session_state.setdefault(seg_key, row.get("seguradora", ""))
        st.session_state.setdefault(sus_key, row.get("susep", ""))
        st.session_state.setdefault(pct_key, row.get("pct", ""))
        st.session_state.setdefault(lmi_key, row.get("lmi", "R$ "))

        seg = c1.text_input("", key=seg_key)
        sus = c2.text_input("", key=sus_key)
        c3.text_input("", key=pct_key, on_change=format_percent_field, args=(pct_key,))
        c4.text_input("", key=lmi_key, on_change=format_money_field, args=(lmi_key,))

        row["seguradora"] = seg
        row["susep"] = sus
        row["pct"] = st.session_state[pct_key]
        row["lmi"] = st.session_state[lmi_key]

        soma_pct += parse_percent(row["pct"])
        soma_lmi += parse_brl_number(row["lmi"])

    # ✅ Totais (Percentual + Total LMI abaixo da coluna LMI – R$)
    st.markdown(f"**Total informado (%):** {fmt_percent_br(soma_pct)} (no Word a linha Total permanece 100%).")

    t1, t2, t3, t4 = st.columns([2.5, 1.2, 1.2, 1.4])
    t1.write("")
    t2.write("")
    t3.markdown("**TOTAL LMI:**")
    t4.text_input("", value=fmt_brl_money(soma_lmi), disabled=True, key="cos_total_lmi_view")

# -------- Limpar / Gerar / Baixar --------
st.markdown("---")
cL, cG = st.columns([1, 3])

with cL:
    if st.button("🧹 Limpar dados", key="btn_limpar_dados"):
        limpar_dados()

with cG:
    if st.button("✅ Gerar Word", key="btn_gerar_word"):
        st.session_state.generated_docx_bytes = build_docx_bytes()
        st.toast("Word gerado com sucesso ✅", icon="✅")

if st.session_state.get("generated_docx_bytes"):
    st.download_button(
        "⬇️ Baixar RN preenchido",
        data=st.session_state.generated_docx_bytes,
        file_name="RN_preenchido.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
